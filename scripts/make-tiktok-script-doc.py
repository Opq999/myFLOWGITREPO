#!/usr/bin/env python3
"""Build a ready-to-read TikTok script as a Word (.docx) document.

Reusable each day: edit the SCRIPT dict (or generate it) and run.
Output goes to docs/scripts/<date>-<slug>.docx
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----- today's pick (random from opqai.solutions) -----
SCRIPT = {
    "day": "Friday 27 June 2026",
    "title": "Explain Your Job to a 5-Year-Old with ChatGPT",
    "url": "https://opqai.solutions/workflows/explain-your-complex-job-simply-using-chatgpt",
    "meta": "~7 min  ·  Beginner  ·  Free (ChatGPT)  ·  Score 10/10",
    "why": "Easiest one on the site, funny, relatable, works on your phone, no spend.",
    "slug": "explain-your-complex-job-simply-using-chatgpt",
    "beats": [
        ("HOOK  (0–3s)", "On-screen text: \"Your mum STILL doesn't know what you do 😭\"",
         "Your family still has no idea what your job actually is? One free prompt fixes that."),
        ("OPEN  (3–8s)", "Open ChatGPT on your phone, start screen-recording here.",
         "Open chatgpt.com — totally free — and start a new chat."),
        ("THE PROMPT  (8–35s)", "SHOW THE PROMPT BIG on screen.",
         "Type this, with your real job title in the bracket:\n\nExplain what a [your job title] does, but make it simple enough for a 5-year-old to understand."),
        ("REACT  (35–55s)", "Read the AI's answer out loud and react.",
         "And look — it says I basically 'count toys and make colourful pictures.' That's... actually accurate 😂"),
        ("PAYOFF  (55–70s)", "Show you can refine it.",
         "Want it simpler or funnier? Just say 'make it even simpler' and it redoes it instantly."),
        ("CTA  (70–75s)", "On-screen text: \"opqai.solutions — link in bio\"",
         "Full steps and the exact prompt are free on opqai.solutions — new AI trick every day. Link in bio."),
    ],
    "caption": "Make AI explain your job like you're 5 😭 Free prompt, link in bio 👇",
    "hashtags": "#ai #chatgpt #howto #naijatech #techtok #productivity",
}


def shade(p, rgb):
    for r in p.runs:
        r.font.color.rgb = RGBColor(*rgb)


def build(s):
    doc = Document()

    # styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # header
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TIKTOK SCRIPT")
    r.bold = True
    r.font.size = Pt(14)
    shade(t, (0x99, 0x00, 0x55))

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run(s["day"]).italic = True

    doc.add_paragraph()

    h = doc.add_heading(s["title"], level=1)
    p = doc.add_paragraph()
    p.add_run(s["meta"]).bold = True
    lp = doc.add_paragraph()
    lp.add_run("Workflow: ").bold = True
    lp.add_run(s["url"])
    wp = doc.add_paragraph()
    wp.add_run("Why this one: ").bold = True
    wp.add_run(s["why"])

    doc.add_paragraph()
    sh = doc.add_heading("🎬 Read this on camera (~75 sec)", level=2)

    for beat, direction, line in s["beats"]:
        bp = doc.add_paragraph()
        br = bp.add_run(beat)
        br.bold = True
        br.font.size = Pt(12)
        shade(bp, (0x99, 0x00, 0x55))

        dp = doc.add_paragraph()
        dr = dp.add_run("[" + direction + "]")
        dr.italic = True
        shade(dp, (0x66, 0x66, 0x66))

        for para in line.split("\n\n"):
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.25)
            run = sp.add_run("“" + para + "”" if not para.startswith("Explain") and "[your job title]" not in para else para)
            # keep the prompt as a clean monospace-ish block
            if "[your job title]" in para or para.startswith("Explain"):
                run.font.name = "Consolas"
                run.bold = True
        doc.add_paragraph()

    doc.add_heading("📝 Caption", level=2)
    doc.add_paragraph(s["caption"])
    doc.add_paragraph(s["hashtags"])

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("After posting: paste the TikTok link into this workflow's tiktokUrl and log it.")
    fr.italic = True
    shade(foot, (0x66, 0x66, 0x66))

    out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs", "scripts")
    os.makedirs(out_dir, exist_ok=True)
    fname = f"{date.today().isoformat()}-{s['slug']}.docx"
    path = os.path.join(out_dir, fname)
    doc.save(path)
    print(path)


if __name__ == "__main__":
    build(SCRIPT)
